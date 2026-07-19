# -*- coding: utf-8 -*-
import re, sys, io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = r"C:\Users\Fung N2\AWS\計畫書_修訂版.md"
OUT = r"C:\Users\Fung N2\AWS\計畫書_修訂版.docx"
CJK = "Microsoft JhengHei"

doc = Document()

# default font (latin + eastasia)
style = doc.styles['Normal']
style.font.name = CJK
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

def set_cjk(run):
    run.font.name = CJK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), CJK)

def add_runs(paragraph, text, base_bold=False):
    # split by **bold**
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for p in parts:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**'):
            r = paragraph.add_run(p[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(p); r.bold = base_bold
        set_cjk(r)

def shade(paragraph, color="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), color)
    pPr.append(shd)

with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
n = len(lines)
while i < n:
    line = lines[i].rstrip('\n')
    s = line.strip()

    # blank
    if s == '':
        i += 1; continue

    # horizontal rule
    if re.fullmatch(r'-{3,}', s):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
        pbdr.append(bottom); pPr.append(pbdr)
        i += 1; continue

    # table block
    if s.startswith('|'):
        block = []
        while i < n and lines[i].strip().startswith('|'):
            block.append(lines[i].strip())
            i += 1
        # parse
        def cells(row):
            return [c.strip() for c in row.strip('|').split('|')]
        header = cells(block[0])
        data_rows = [cells(r) for r in block[2:]]  # skip separator
        cols = len(header)
        tbl = doc.add_table(rows=1, cols=cols)
        tbl.style = 'Light Grid Accent 1'
        for j, h in enumerate(header):
            cell = tbl.rows[0].cells[j]
            cell.paragraphs[0].text = ''
            add_runs(cell.paragraphs[0], h, base_bold=True)
        for dr in data_rows:
            row = tbl.add_row().cells
            for j in range(cols):
                txt = dr[j] if j < len(dr) else ''
                row[j].paragraphs[0].text = ''
                add_runs(row[j].paragraphs[0], txt)
        doc.add_paragraph()
        continue

    # headings
    if s.startswith('#### '):
        p = doc.add_heading(level=3); add_runs(p, s[5:]); 
        for r in p.runs: set_cjk(r)
        i += 1; continue
    if s.startswith('### '):
        p = doc.add_heading(level=2); add_runs(p, s[4:])
        for r in p.runs: set_cjk(r)
        i += 1; continue
    if s.startswith('## '):
        p = doc.add_heading(level=1); add_runs(p, s[3:])
        for r in p.runs: set_cjk(r)
        i += 1; continue
    if s.startswith('# '):
        p = doc.add_heading(level=0); add_runs(p, s[2:])
        for r in p.runs: set_cjk(r)
        i += 1; continue

    # blockquote (may span multiple lines)
    if s.startswith('>'):
        while i < n and lines[i].strip().startswith('>'):
            q = lines[i].strip()[1:].strip()
            if q == '':
                i += 1; continue
            if q.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, q[2:])
            else:
                p = doc.add_paragraph()
                add_runs(p, q)
            p.paragraph_format.left_indent = Pt(18)
            shade(p, "EEF3FA")
            i += 1
        continue

    # bullet (with nesting)
    m = re.match(r'^(\s*)- (.*)$', line)
    if m:
        indent = len(m.group(1))
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, m.group(2))
        if indent >= 2:
            p.paragraph_format.left_indent = Pt(36)
        i += 1; continue

    # numbered list
    m = re.match(r'^(\s*)\d+\. (.*)$', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        add_runs(p, m.group(2))
        i += 1; continue

    # normal paragraph
    p = doc.add_paragraph()
    add_runs(p, s)
    i += 1

doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
